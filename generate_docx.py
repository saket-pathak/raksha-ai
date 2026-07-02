import os
import sys
from pathlib import Path
import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Helpers for XML manipulation in python-docx
def set_cell_background(cell, color_hex):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:val'), 'clear')
    shading_elm.set(qn('w:color'), 'auto')
    shading_elm.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_table_borders(table):
    tblPr = table._tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for b in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{b}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4') # 0.5 pt
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'D1D5DB') # Light gray border
        tblBorders.append(border)
    tblPr.append(tblBorders)

def add_heading(doc, text, level):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level > 1 else 28)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    
    run = p.add_run(text)
    run.font.name = 'Segoe UI'
    if level == 1:
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(15, 32, 67) # Deep Blue/Navy
        run.font.bold = True
    elif level == 2:
        run.font.size = Pt(15)
        run.font.color.rgb = RGBColor(26, 54, 93) # Dark Navy
        run.font.bold = True
    elif level == 3:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(74, 85, 104) # Slate Blue
        run.font.bold = True
    return p

def add_p(doc, text, bold_prefix=None, space_after=6, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.bold = True
        run_b.font.name = 'Calibri'
        run_b.font.size = Pt(10.5)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(10.5)
    run.font.italic = italic
    run.font.color.rgb = RGBColor(51, 65, 85) # Slate 700
    return p

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.bold = True
        run_b.font.name = 'Calibri'
        run_b.font.size = Pt(10.5)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(51, 65, 85)
    return p

def add_code_block(doc, code_text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    
    set_cell_background(cell, 'F9FAFB') # Light gray bg
    set_cell_margins(cell, top=120, bottom=120, left=180, right=180) # 6pt top/bottom, 9pt left/right
    
    # Border
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    
    # Left thick accent border
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '24') # 3pt width
    left.set(qn('w:space'), '0')
    left.set(qn('w:color'), '3B82F6') # Blue accent
    tcBorders.append(left)
    
    # Other thin borders
    for border_name in ['top', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4') # 0.5 pt
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'E5E7EB')
        tcBorders.append(border)
    tcPr.append(tcBorders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(31, 41, 55)

def add_packages_table(doc, packages_data):
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Headers
    hdr_cells = table.rows[0].cells
    headers = ["Package Name", "Version Scope", "Core Responsibility in Ecosystem", "Environment"]
    widths = [Pt(110), Pt(70), Pt(220), Pt(80)]
    
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        hdr_cells[i].width = widths[i]
        
        # Color headers
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1A365D') # Deep Navy
        hdr_cells[i]._tc.get_or_add_tcPr().append(shd)
        
        # Font styling for header
        p = hdr_cells[i].paragraphs[0]
        run = p.runs[0]
        run.font.name = 'Segoe UI'
        run.font.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=150, right=150)
        
    for name, version, desc, env in packages_data:
        row_cells = table.add_row().cells
        data = [name, version, desc, env]
        for i in range(4):
            row_cells[i].text = data[i]
            row_cells[i].width = widths[i]
            p = row_cells[i].paragraphs[0]
            run = p.runs[0]
            run.font.name = 'Calibri'
            run.font.size = Pt(9.5)
            
            # Alternate row background
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:val'), 'clear')
            shading_elm.set(qn('w:color'), 'auto')
            shading_elm.set(qn('w:fill'), 'F9FAFB')
            row_cells[i]._tc.get_or_add_tcPr().append(shading_elm)
            set_cell_margins(row_cells[i], top=100, bottom=100, left=150, right=150)

def build_document():
    doc = docx.Document()
    
    # Margins setup
    section = doc.sections[0]
    section.top_margin = Pt(72)
    section.bottom_margin = Pt(72)
    section.left_margin = Pt(72)
    section.right_margin = Pt(72)
    
    # ----------------- COVER PAGE -----------------
    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_before = Pt(120)
    
    p_title = doc.add_paragraph()
    run_title = p_title.add_run("🛡️ RAKSHA AI")
    run_title.font.name = 'Segoe UI'
    run_title.font.size = Pt(36)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(15, 32, 67) # Slate 800
    p_title.paragraph_format.space_after = Pt(6)
    
    p_subtitle = doc.add_paragraph()
    run_subtitle = p_subtitle.add_run("Intelligent Road Safety Ecosystem")
    run_subtitle.font.name = 'Segoe UI'
    run_subtitle.font.size = Pt(18)
    run_subtitle.font.color.rgb = RGBColor(71, 85, 105) # Slate 600
    p_subtitle.paragraph_format.space_after = Pt(220)
    
    p_meta = doc.add_paragraph()
    p_meta.paragraph_format.line_spacing = 1.3
    run_meta1 = p_meta.add_run("Technical Architecture, Design Assumptions & Core Codebase Documentation\n")
    run_meta1.font.name = 'Calibri'
    run_meta1.font.bold = True
    run_meta1.font.size = Pt(11)
    run_meta1.font.color.rgb = RGBColor(15, 32, 67)
    
    run_meta2 = p_meta.add_run("Version: 2.0.0 (Stable Release)\n")
    run_meta2.font.name = 'Calibri'
    run_meta2.font.size = Pt(10)
    run_meta2.font.color.rgb = RGBColor(100, 116, 139)
    
    run_meta3 = p_meta.add_run("Date: May 31, 2026\n")
    run_meta3.font.name = 'Calibri'
    run_meta3.font.size = Pt(10)
    run_meta3.font.color.rgb = RGBColor(100, 116, 139)
    
    run_meta4 = p_meta.add_run("Maintainer & Architect: Saket Pathak\n")
    run_meta4.font.name = 'Calibri'
    run_meta4.font.size = Pt(10)
    run_meta4.font.color.rgb = RGBColor(100, 116, 139)
    
    doc.add_page_break()
    
    # ----------------- SECTION 1: INTRODUCTION -----------------
    add_heading(doc, "1. Executive Summary & Project Introduction", 1)
    
    add_p(doc, "India records one of the highest numbers of road accidents globally, resulting in tens of thousands of preventable fatalities annually. The critical challenges that exacerbate this issue include delayed emergency medical responses, lack of active hazard awareness, poor monitoring of structural road defects, and linguistic barriers that restrict citizens from reporting hazards to local authorities.")
    add_p(doc, "Raksha AI is designed as a unified, intelligent road safety ecosystem. It bridges the gap between active road hazards, citizen reporting, real-time safety navigation, and administrative response through three primary pillars:")
    
    add_bullet(doc, " empowering citizens to report infrastructure issues in their native languages with direct geo-location tags.", "Crowdsourced Localized Incident Reporting:")
    add_bullet(doc, " using automated computer-vision algorithms (Gemini Vision API integration) to identify and classify defects like potholes, damaged roads, and waterlogging from uploaded images.", "Intelligent Visual Defect Classification:")
    add_bullet(doc, " providing real-time risk scores for routes based on geographical, weather, traffic, and temporal inputs.", "Active Routing and Risk Profiling:")
    add_bullet(doc, " offering a single-tap SOS protocol that uses network location resolution, notifies emergency contacts, and maps the closest medical facilities.", "Emergency Resiliency Framework:")
    
    # ----------------- SECTION 2: SYSTEM ARCHITECTURE -----------------
    add_heading(doc, "2. Architectural Design & System Components", 1)
    
    add_p(doc, "Raksha AI utilizes a decoupled client-server architecture with a clear separation between the UI layer, application logic, and the AI evaluation pipelines. This design enables independent scaling, quick localized modifications, and seamless integration with external third-party systems.")
    
    add_heading(doc, "2.1 Frontend Component Layer", 2)
    add_p(doc, "The frontend is built as a single-page application (SPA) using React and Vite. It contains the following modules:")
    add_bullet(doc, " Provides a command-center interface displaying active incident maps (Leaflet overlay), recent hazards, hotspot distributions, and active alert counters.", "Tactical Dashboard:")
    add_bullet(doc, " Houses the form to capture details, drag-and-drop image uploads, and handles communication with the backend detection endpoints.", "Report & Detection UI:")
    add_bullet(doc, " Handles the coordinates capturing, emergency trigger countdowns, and displays nearest hospital cards.", "SOS Interface:")
    add_bullet(doc, " Encapsulates the LanguageProvider state context, allowing instant translation switching across all pages.", "Multilingual Core:")
    
    add_heading(doc, "2.2 Backend Application Layer", 2)
    add_p(doc, "The backend is implemented using Flask (Python) and serves as the API gateway. It manages all application business logic and integrates with helper services:")
    add_bullet(doc, " Configures environment and server variables (upload directory, APIs keys).", "Config Service:")
    add_bullet(doc, " Encapsulates the AI evaluation pipeline and coordinates requests to the AI models.", "AI Bridge:")
    add_bullet(doc, " Implements JWT token extraction, session login/logout, and credentials verification.", "Authentication Service:")
    add_bullet(doc, " A thread-safe file database service that reads and writes reported issues to local JSON storage.", "Reports Service:")
    add_bullet(doc, " Handles IP-based geocoding calculations and calculates distances to nearest medical structures.", "Maps Service:")
    
    add_heading(doc, "2.3 AI & Evaluation Pipelines", 2)
    add_p(doc, "The ecosystem operates two distinct AI/ML components:")
    add_bullet(doc, " Accessible via `/roads/detect`, it handles image uploads. It first attempts to leverage the Google Gemini Vision API (using the 'gemini-2.5-flash' model) to analyze the image details. If the API is unavailable or unconfigured, it utilizes a filename-heuristic algorithm and brightness/contrast analysis to output deterministic mock metrics.", "Road Defect Classification Model:")
    add_bullet(doc, " Accessible via `/risk/score` and `/risk/route-profile`, it computes safety risk indices (from 0 to 99) based on factors like time of day, weather conditions, traffic level, road quality, and specific geographical features.", "Tabular Risk Scoring Engine:")
    
    doc.add_page_break()
    
    # ----------------- SECTION 3: KEY SYSTEM FEATURES -----------------
    add_heading(doc, "3. Key Features & Operational Workflows", 1)
    
    add_heading(doc, "3.1 Smart SOS Emergency System", 2)
    add_p(doc, "When a citizen triggers an SOS, the system executes the following steps:")
    add_bullet(doc, " Resolves the user's location coordinates. It first requests the browser geolocation; if unavailable, it uses the backend's IP Geolocation service (`ipinfo.io`) as a fallback.")
    add_bullet(doc, " Calculates distances to nearest hospitals using coordinates math.")
    add_bullet(doc, " Notifies emergency contacts and writes the detailed incident log to 'sos_logs.txt' on the server.")
    add_bullet(doc, " Streams the dispatched SOS alert to the active command-center dashboard.")
    
    add_heading(doc, "3.2 AI Road Issue Detection", 2)
    add_p(doc, "Citizens snap a picture of a road hazard and upload it. The pipeline:")
    add_bullet(doc, " Validates file extensions (PNG, JPG, JPEG, WEBP, HEIC) and saves the image to the uploads directory.")
    add_bullet(doc, " Resizes the image to 224x224 and converts it to RGB to prepare it for processing.")
    add_bullet(doc, " Calls the Gemini Vision API using a highly structured prompt to identify potholes, waterlogging, or general damage.")
    add_bullet(doc, " If Gemini fails, it checks the filename for keywords (e.g., 'pothole', 'water') and computes a deterministic confidence score based on the image's standard deviation and mean pixel brightness.")
    
    add_heading(doc, "3.3 Route Risk Profiler", 2)
    add_p(doc, "The risk model accepts a list of path waypoints and calculates safety hazard indices:")
    add_bullet(doc, " Evaluates risk weights: time of day (e.g. peak rush hours), weather conditions (e.g. rain/fog), road reports, and traffic levels.")
    add_bullet(doc, " Applies geographical modifiers (e.g. flyovers, junctions, and specific districts add incremental risk).")
    add_bullet(doc, " Streams active risk alerts to connected frontend clients using a Server-Sent Events (SSE) `/risk/stream` protocol.")
    
    add_heading(doc, "3.4 Multilingual Core (i18n)", 2)
    add_p(doc, "To bridge language barriers, Raksha AI implements full dynamic translation. Localization is supported for six languages: English (en), Hindi (hi), Tamil (ta), Telugu (te), Kannada (kn), and Malayalam (ml).")
    add_bullet(doc, " Frontend localization uses react-i18next and a custom state context provider. User preferences are persisted in local storage.")
    add_bullet(doc, " Backend localization dynamically translates error messages, validation statuses, and notifications based on the query parameter `?language=xx` or the `Accept-Language` HTTP header.")
    
    doc.add_page_break()
    
    # ----------------- SECTION 4: SOFTWARE PACKAGES USED -----------------
    add_heading(doc, "4. Software Requirements & Packages Used", 1)
    add_p(doc, "The following table lists the core software libraries and frameworks utilized in the development and runtime execution of the Raksha AI ecosystem:")
    
    packages_data = [
        ("Flask", ">=3.0, <4.0", "Core WSGI Python web application framework for hosting endpoints and middleware.", "Backend"),
        ("flask-cors", ">=4.0, <5.0", "Cross-Origin Resource Sharing middleware for Flask, facilitating secure web client requests.", "Backend"),
        ("requests", ">=2.31, <3.0", "HTTP client library for retrieving IP geolocation data and external maps integration.", "Backend"),
        ("Pillow", ">=10.0, <11.0", "Python Imaging Library for image validation, metadata extraction, and preprocessing.", "Backend"),
        ("numpy", ">=1.26", "Numeric math library used for image array operations and brightness calculations.", "Backend"),
        ("firebase-admin", ">=6.5, <7.0", "Firebase SDK for managing administrative operations, authentication tokens, and synchronizations.", "Backend"),
        ("python-dotenv", ">=1.0, <2.0", "Parses and loads environment variables from a .env file into the system runtime.", "Backend"),
        ("google-genai", ">=0.1.0", "Google Gemini SDK to interface with Gemini models (e.g. gemini-2.5-flash) for vision and AI tasks.", "Backend"),
        ("react", "^18.3.1", "Component-based Javascript UI library for managing single-page application rendering.", "Frontend"),
        ("react-dom", "^18.3.1", "Serves as the entry point to the DOM, coupling React with browser rendering.", "Frontend"),
        ("react-router-dom", "^6.28.0", "Client-side routing engine to handle page switching and navigation parameters.", "Frontend"),
        ("leaflet", "^1.9.4", "Javascript interactive map library for rendering hotspot locations and route hazard layers.", "Frontend"),
        ("i18next", "^23.7.6", "Core internationalization and translation framework.", "Frontend"),
        ("react-i18next", "^14.0.0", "React components and hooks wrapper for the i18next localization library.", "Frontend"),
        ("i18next-browser-languagedetector", "^7.2.0", "Detects preferred user languages from browser settings automatically.", "Frontend"),
        ("i18next-http-backend", "^2.4.2", "Asynchronously fetches translations from backend directories or configuration objects.", "Frontend"),
        ("vite", "^5.4.10", "High-performance frontend build tool and hot-module development server.", "Frontend Dev")
    ]
    
    add_packages_table(doc, packages_data)
    add_table_borders(doc.tables[0])
    
    # ----------------- SECTION 5: TECHNICAL ASSUMPTIONS -----------------
    add_heading(doc, "5. Technical Assumptions & Design Decisions", 1)
    
    add_bullet(doc, " The backend requires external network connectivity to reach the Google Gemini API (for vision classification) and the IP Geolocation API (for fallback SOS coordinates). If these networks are offline, the application falls back gracefully to local calculations and pre-configured coordinates without crashing.", "API Dependencies & Fallbacks:")
    add_bullet(doc, " The development environment utilizes a thread-safe local JSON database file (reports.json) and local logs (sos_logs.txt) to avoid database setup overhead. This design is highly portable and easily migrates to Firebase or SQL systems by changing the service implementation.", "In-Memory & Local Storage:")
    add_bullet(doc, " The AI image detection model requires a valid 'GEMINI_API_KEY' in the environmental settings. In the absence of a key, the system leverages a filename-keyword check and basic image statistics (mean and deviation) to classify the image deterministically, enabling testing with sample image names.", "Model Availability:")
    add_bullet(doc, " Time coordinates, traffic levels, and weather data are supplied by the client during risk requests (often mapped from local sensors or weather services). If missing, the backend defaults to standard daylight and clear conditions.", "Risk Engine Inputs:")
    add_bullet(doc, " The localization service assumes UTF-8 encoding across all configuration and translation files to prevent character corruption of Indian scripts (Hindi, Tamil, Telugu, Kannada, Malayalam).", "Encoding and Fonts:")
    
    doc.add_page_break()
    
    # ----------------- SECTION 6: SOURCE CODE -----------------
    add_heading(doc, "6. Complete Source Code Repository", 1)
    add_p(doc, "This section contains the core, functional source code of the Raksha AI backend and testing suite. The code has been structured and documented for reference:")
    
    code_files = [
        ("backend/config.py", "System and Environment Configuration Loader"),
        ("backend/models/RiskModel.py", "Tabular Risk Prediction Scoring Model"),
        ("backend/models/RoadModel.py", "AI Road Issue Detection and Gemini Vision Bridge"),
        ("backend/models/SosModel.py", "SOS Workflow, Geolocation, and Dispatch Simulation"),
        ("backend/services/localization_service.py", "Multilingual Backend Translation Service"),
        ("backend/services/reports_service.py", "Thread-Safe JSON Reports Storage & Management"),
        ("backend/main.py", "Flask Core Router & Main Gateway Server"),
        ("tests/test_backend_routes.py", "Automated Route Integration and Quality Assurance Tests")
    ]
    
    for relative_path, description in code_files:
        add_heading(doc, f"6.X File: {relative_path}", 2)
        add_p(doc, description, italic=True)
        
        file_path = Path(relative_path)
        if file_path.exists():
            try:
                code_text = file_path.read_text(encoding="utf-8")
                add_code_block(doc, code_text)
            except Exception as e:
                add_code_block(doc, f"Error reading file {relative_path}: {e}")
        else:
            add_code_block(doc, f"File not found: {relative_path}")
            
    # Save the output file
    output_filename = "Raksha_AI_System_Documentation.docx"
    doc.save(output_filename)
    print(f"Document successfully created at {output_filename}")

if __name__ == "__main__":
    build_document()
